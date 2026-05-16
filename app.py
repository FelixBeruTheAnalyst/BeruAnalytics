import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import plotly.express as px
from plotly.subplots import make_subplots
import io
import os
import json
from datetime import datetime

# ================================================
# PAGE CONFIG
# ================================================
st.set_page_config(
    page_title="BeruAnalytics",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ================================================
# CUSTOM CSS
# ================================================
st.markdown("""
<style>
    .main { background-color: #F8F9FA; }
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}

    .hero {
        background: linear-gradient(135deg, #1A6B72 0%, #1A3A5C 100%);
        border-radius: 16px;
        padding: 48px 40px;
        color: white;
        margin-bottom: 32px;
    }
    .hero h1 { font-size: 42px; font-weight: 800; margin: 0 0 12px 0; color: white; }
    .hero p { font-size: 18px; opacity: 0.9; margin: 0 0 24px 0; color: white; }

    .feature-card {
        background: white;
        border-radius: 12px;
        padding: 24px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.06);
        height: 100%;
        border-top: 3px solid #1A6B72;
    }
    .feature-card h3 { color: #1A6B72; font-size: 18px; margin-bottom: 8px; }
    .feature-card p { color: #555555; font-size: 14px; line-height: 1.6; }

    .section-header {
        font-size: 22px;
        font-weight: 700;
        color: #1A1A2E;
        border-bottom: 3px solid #1A6B72;
        padding-bottom: 8px;
        margin-bottom: 24px;
    }

    .nav-btn-container {
        display: flex;
        gap: 12px;
        margin: 24px 0;
        flex-wrap: wrap;
    }
    .nav-btn {
        background: white;
        border: 2px solid #1A6B72;
        border-radius: 8px;
        padding: 12px 20px;
        color: #1A6B72;
        font-weight: 600;
        font-size: 14px;
        cursor: pointer;
        transition: all 0.2s ease;
        text-decoration: none;
        display: inline-block;
        position: relative;
    }
    .nav-btn:hover {
        background: #1A6B72;
        color: white;
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(26,107,114,0.3);
    }
    .nav-btn .tooltip {
        visibility: hidden;
        background: #1A1A2E;
        color: white;
        text-align: center;
        border-radius: 6px;
        padding: 6px 10px;
        position: absolute;
        z-index: 1;
        bottom: 125%;
        left: 50%;
        transform: translateX(-50%);
        white-space: nowrap;
        font-size: 12px;
        font-weight: 400;
        opacity: 0;
        transition: opacity 0.2s;
    }
    .nav-btn:hover .tooltip {
        visibility: visible;
        opacity: 1;
    }
</style>
""", unsafe_allow_html=True)

# ================================================
# SESSION STATE
# ================================================
if 'df' not in st.session_state:
    st.session_state.df = None
if 'file_name' not in st.session_state:
    st.session_state.file_name = None
if 'analysis' not in st.session_state:
    st.session_state.analysis = None
if 'page' not in st.session_state:
    st.session_state.page = "🏠 Home"

# ================================================
# SIDEBAR
# ================================================
with st.sidebar:
    st.markdown("""
    <div style='text-align:center; padding: 16px 0;'>
        <h2 style='color:#1A6B72; margin:0; font-size:24px;'>📊 BeruAnalytics</h2>
        <p style='color:#888; font-size:12px; margin:4px 0 0 0;'>AI-Powered Data Analytics</p>
    </div>
    """, unsafe_allow_html=True)

    st.divider()

    pages = ["🏠 Home", "📊 Dashboard", "🤖 BeruDataNarrate",
             "💬 AI Assistant", "📋 Data Explorer", "ℹ️ About"]

    page = st.radio(
        "Navigate",
        pages,
        label_visibility="collapsed",
        index=pages.index(st.session_state.page)
    )
    st.session_state.page = page

    st.divider()

    if st.session_state.df is not None:
        df = st.session_state.df
        st.success("✅ Data loaded")
        st.caption(f"📁 {st.session_state.file_name}")
        st.caption(f"📊 {len(df):,} rows × {len(df.columns)} cols")
        if st.button("🗑️ Clear Data", use_container_width=True):
            st.session_state.df = None
            st.session_state.file_name = None
            st.session_state.analysis = None
            st.rerun()
    else:
        st.warning("⚠️ No data loaded")
        st.caption("Upload a file to get started")

    st.divider()
    st.caption("Built by Felix Beru Tsinzole")
    st.caption("Nairobi, Kenya 🇰🇪")
# ================================================
# HELPER FUNCTIONS
# ================================================
def load_data(uploaded_file):
    try:
        if uploaded_file.name.endswith('.csv'):
            df = pd.read_csv(uploaded_file, encoding='latin1')
        else:
            df = pd.read_excel(uploaded_file)
        df.columns = df.columns.str.strip()

        if len(df) == 0:
            st.error("❌ The uploaded file is empty.")
            return None

        warnings_list = []
        if len(df) > 100000:
            warnings_list.append(f"⚠️ Large dataset ({len(df):,} rows) — analysis may be slower")
        missing = df.isnull().sum().sum()
        if missing > 0:
            pct = (missing / (len(df) * len(df.columns))) * 100
            warnings_list.append(f"⚠️ {missing:,} missing values found ({pct:.1f}%) — may affect accuracy")
        dupes = df.duplicated().sum()
        if dupes > 0:
            warnings_list.append(f"⚠️ {dupes:,} duplicate rows detected")
        if len(df.columns) < 2:
            warnings_list.append("⚠️ Only 1 column detected — charts work best with multiple columns")
        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        if len(numeric_cols) == 0:
            warnings_list.append("⚠️ No numeric columns — some charts may not be available")

        for w in warnings_list:
            st.warning(w)

        return df
    except Exception as e:
        st.error(f"❌ Could not read this file. Please upload a valid CSV or Excel file. Details: {e}")
        return None


def get_sample_data():
    np.random.seed(42)
    counties = ['Nairobi', 'Mombasa', 'Kisumu', 'Nakuru', 'Eldoret',
                'Thika', 'Malindi', 'Kitale', 'Garissa', 'Nyeri']
    products = ['Solar Panel', 'Battery', 'Inverter', 'Cable', 'Controller']
    months = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun',
              'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
    data = []
    for _ in range(200):
        data.append({
            'Month': np.random.choice(months),
            'County': np.random.choice(counties),
            'Product': np.random.choice(products),
            'Units_Sold': np.random.randint(10, 500),
            'Revenue_KES': np.random.randint(5000, 200000),
            'Cost_KES': np.random.randint(3000, 150000),
            'Customer_Rating': round(np.random.uniform(3.0, 5.0), 1),
            'Returns': np.random.randint(0, 20),
        })
    df = pd.DataFrame(data)
    df['Profit_KES'] = df['Revenue_KES'] - df['Cost_KES']
    return df


def get_groq_key():
    try:
        return st.secrets["GROQ_API_KEY"]
    except:
        return None


# ================================================
# PAGE 1 — HOME
# ================================================
if page == "🏠 Home":

    st.markdown("""
    <div class='hero'>
        <h1>📊 BeruAnalytics</h1>
        <p>Upload your data. Get instant AI-powered analysis,
        interactive dashboards and professional reports in minutes.</p>
        <p style='font-size:14px; opacity:0.7;'>Powered by AI | Built for Kenya and beyond 🇰🇪</p>
    </div>
    """, unsafe_allow_html=True)

    # ── FUNCTIONAL NAVIGATION BUTTONS ──
    st.markdown("### 🧭 Quick Navigation")
    st.markdown("*Click any feature to go directly to it*")

    nav1, nav2, nav3, nav4, nav5 = st.columns(5)

    with nav1:
        st.markdown("""
        <div class='feature-card' style='text-align:center; cursor:pointer;'>
            <h3 style='font-size:24px; margin:0;'>📊</h3>
            <h3>Dashboard</h3>
            <p>Interactive charts and KPI cards — bar, pie, scatter,
            heatmap and line charts generated instantly from your data.</p>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Go to Dashboard →",
                      key="nav_dash",
                      use_container_width=True,
                      type="primary"):
            st.session_state.page = "📊 Dashboard"
            st.rerun()

    with nav2:
        st.markdown("""
        <div class='feature-card' style='text-align:center; cursor:pointer;'>
            <h3 style='font-size:24px; margin:0;'>🤖</h3>
            <h3>BeruDataNarrate</h3>
            <p>AI-powered report generator — upload data, get a
            professional Word and PDF report in minutes.</p>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Go to Reports →",
                      key="nav_narrate",
                      use_container_width=True,
                      type="primary"):
            st.session_state.page = "🤖 BeruDataNarrate"
            st.rerun()

    with nav3:
        st.markdown("""
        <div class='feature-card' style='text-align:center; cursor:pointer;'>
            <h3 style='font-size:24px; margin:0;'>💬</h3>
            <h3>AI Assistant</h3>
            <p>Ask questions about your data in plain English and
            get instant AI-powered answers and insights.</p>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Go to Assistant →",
                      key="nav_chat",
                      use_container_width=True,
                      type="primary"):
            st.session_state.page = "💬 AI Assistant"
            st.rerun()

    with nav4:
        st.markdown("""
        <div class='feature-card' style='text-align:center; cursor:pointer;'>
            <h3 style='font-size:24px; margin:0;'>📋</h3>
            <h3>Data Explorer</h3>
            <p>Search, filter, sort and export your data as
            CSV or Excel with one click.</p>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Go to Explorer →",
                      key="nav_explore",
                      use_container_width=True,
                      type="primary"):
            st.session_state.page = "📋 Data Explorer"
            st.rerun()

    with nav5:
        st.markdown("""
        <div class='feature-card' style='text-align:center; cursor:pointer;'>
            <h3 style='font-size:24px; margin:0;'>ℹ️</h3>
            <h3>About</h3>
            <p>Privacy policy, data security information
            and platform details.</p>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Go to About →",
                      key="nav_about",
                      use_container_width=True,
                      type="primary"):
            st.session_state.page = "ℹ️ About"
            st.rerun()

    # Get Started
    st.markdown("<div class='section-header'>🚀 Get Started</div>",
                unsafe_allow_html=True)

    col_upload, col_sample = st.columns([3, 1])

    with col_upload:
        uploaded = st.file_uploader(
            "Upload your CSV or Excel file to begin",
            type=['csv', 'xlsx', 'xls'],
            help="Supported: CSV, Excel | Max recommended: 100,000 rows"
        )

    with col_sample:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("🎯 Try Sample Data",
                      use_container_width=True,
                      help="Load a sample Kenya sales dataset to explore all features"):
            df_sample = get_sample_data()
            st.session_state.df = df_sample
            st.session_state.file_name = "sample_kenya_sales_data.csv"
            st.success("✅ Sample dataset loaded — 200 rows of Kenya sales data!")
            st.info("👈 Navigate to Dashboard or BeruDataNarrate to explore")

    if uploaded:
        df = load_data(uploaded)
        if df is not None:
            st.session_state.df = df
            st.session_state.file_name = uploaded.name

            numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()

            st.divider()
            st.markdown("### ✅ Data Quality Report")

            q1, q2, q3, q4, q5 = st.columns(5)
            with q1:
                st.metric("Total Rows", f"{len(df):,}")
            with q2:
                st.metric("Total Columns", len(df.columns))
            with q3:
                missing = df.isnull().sum().sum()
                st.metric("Missing Values", missing,
                          delta="Clean ✅" if missing == 0 else "⚠️ Found",
                          delta_color="normal" if missing == 0 else "inverse")
            with q4:
                dupes = df.duplicated().sum()
                st.metric("Duplicates", dupes,
                          delta="Clean ✅" if dupes == 0 else "⚠️ Found",
                          delta_color="normal" if dupes == 0 else "inverse")
            with q5:
                st.metric("Numeric Cols", len(numeric_cols))

            st.success(f"✅ **{uploaded.name}** loaded — {len(df):,} rows, {len(df.columns)} columns")
            st.info("👈 Use the sidebar to navigate to Dashboard, BeruDataNarrate or Data Explorer")

            with st.expander("👀 Preview first 5 rows"):
                st.dataframe(df.head(), use_container_width=True)

    elif st.session_state.df is not None:
        st.success(f"✅ **{st.session_state.file_name}** is loaded and ready")
        st.info("👈 Navigate using the sidebar")

    st.divider()

    # Feature cards
    st.markdown("<div class='section-header'>✨ What BeruAnalytics Does</div>",
                unsafe_allow_html=True)

    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("""<div class='feature-card'>
            <h3>📊 Interactive Dashboard</h3>
            <p>Automatic KPI cards, bar charts, line charts, scatter plots
            and correlation heatmaps generated instantly from your data.</p>
        </div>""", unsafe_allow_html=True)
    with col2:
        st.markdown("""<div class='feature-card'>
            <h3>🤖 AI Report Generator</h3>
            <p>BeruDataNarrate analyzes your data with AI and generates
            a professional Word and PDF report with executive summary,
            key findings and recommendations.</p>
        </div>""", unsafe_allow_html=True)
    with col3:
        st.markdown("""<div class='feature-card'>
            <h3>💬 AI Data Assistant</h3>
            <p>Ask questions about your data in plain English.
            Get instant AI-powered answers and insights
            without writing any code.</p>
        </div>""", unsafe_allow_html=True)

    st.divider()

    col4, col5, col6 = st.columns(3)
    with col4:
        st.markdown("""<div class='feature-card'>
            <h3>📋 Data Explorer</h3>
            <p>Browse, search, sort and filter your data in a
            smart interactive table. Export as CSV or Excel instantly.</p>
        </div>""", unsafe_allow_html=True)
    with col5:
        st.markdown("""<div class='feature-card'>
            <h3>🔒 Data Quality Checks</h3>
            <p>Automatic checks for missing values, duplicates and
            data integrity issues on every upload — so your analysis
            is always based on clean data.</p>
        </div>""", unsafe_allow_html=True)
    with col6:
        st.markdown("""<div class='feature-card'>
            <h3>📥 Export Everything</h3>
            <p>Download your analysis as Word documents, PDF reports,
            CSV or Excel files. Share with stakeholders instantly.</p>
        </div>""", unsafe_allow_html=True)

    st.divider()

    s1, s2, s3, s4 = st.columns(4)
    with s1:
        st.metric("File Formats", "CSV + Excel", delta="Supported")
    with s2:
        st.metric("Chart Types", "6+", delta="Interactive")
    with s3:
        st.metric("Report Formats", "Word + PDF", delta="Professional")
    with s4:
        st.metric("AI Powered", "Yes", delta="Groq LLaMA")

    st.divider()
    st.caption("BeruAnalytics | Built by Felix Beru Tsinzole | Nairobi, Kenya 🇰🇪")

# ================================================
# PAGE 2 — DASHBOARD
# ================================================
elif page == "📊 Dashboard":
    st.markdown("<div class='section-header'>📊 Interactive Dashboard</div>",
                unsafe_allow_html=True)

    if st.session_state.df is None:
        st.warning("⚠️ No data loaded. Please upload a file on the Home page first.")
        st.stop()

    df = st.session_state.df
    numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    categorical_cols = df.select_dtypes(include=['object']).columns.tolist()

    with st.sidebar:
        st.markdown("### 🔍 Filters")
        filtered_df = df.copy()
        for col in numeric_cols[:3]:
            min_val = float(df[col].min())
            max_val = float(df[col].max())
            if min_val != max_val:
                selected = st.slider(f"{col}", min_val, max_val, (min_val, max_val))
                filtered_df = filtered_df[
                    (filtered_df[col] >= selected[0]) &
                    (filtered_df[col] <= selected[1])
                ]
        for col in categorical_cols[:2]:
            options = df[col].unique().tolist()
            selected_cats = st.multiselect(f"{col}", options, default=options)
            if selected_cats:
                filtered_df = filtered_df[filtered_df[col].isin(selected_cats)]
        st.caption(f"Showing {len(filtered_df):,} of {len(df):,} rows")

    st.subheader("📌 Key Metrics")
    kpi_cols = st.columns(len(numeric_cols[:4]) if numeric_cols else 1)
    for i, col in enumerate(numeric_cols[:4]):
        with kpi_cols[i]:
            total = filtered_df[col].sum()
            mean = filtered_df[col].mean()
            st.metric(label=col, value=f"{mean:.1f}", delta=f"Total: {total:,.1f}")

    st.divider()

    if numeric_cols and categorical_cols:
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("📊 Bar Chart")
            x_col = st.selectbox("X Axis (Category)", categorical_cols, key="bar_x")
            y_col = st.selectbox("Y Axis (Value)", numeric_cols, key="bar_y")
            df_bar = filtered_df.groupby(x_col)[y_col].mean().reset_index()
            df_bar = df_bar.sort_values(y_col, ascending=True)
            fig_bar = go.Figure(go.Bar(
                x=df_bar[y_col], y=df_bar[x_col], orientation='h',
                marker_color='#1A6B72',
                hovertemplate=f'<b>%{{y}}</b><br>{y_col}: %{{x:.2f}}<extra></extra>'
            ))
            fig_bar.update_layout(height=400, plot_bgcolor='#f8f9fa',
                                   paper_bgcolor='white', xaxis_title=y_col,
                                   yaxis=dict(tickfont=dict(size=10)),
                                   margin=dict(l=0, r=20, t=20, b=0))
            st.plotly_chart(fig_bar, use_container_width=True)

        with col2:
            st.subheader("🥧 Distribution")
            pie_col = st.selectbox("Category Column", categorical_cols, key="pie_col")
            pie_val = st.selectbox("Value Column", numeric_cols, key="pie_val")
            df_pie = filtered_df.groupby(pie_col)[pie_val].sum().reset_index()
            fig_pie = go.Figure(go.Pie(
                labels=df_pie[pie_col], values=df_pie[pie_val], hole=0.4,
                hovertemplate='<b>%{label}</b><br>Value: %{value:.2f}<br>Share: %{percent}<extra></extra>'
            ))
            fig_pie.update_layout(height=400, paper_bgcolor='white',
                                   margin=dict(l=0, r=0, t=20, b=0))
            st.plotly_chart(fig_pie, use_container_width=True)

    st.divider()

    if len(numeric_cols) >= 2:
        col3, col4 = st.columns(2)
        with col3:
            st.subheader("📈 Scatter Plot")
            x_scatter = st.selectbox("X Axis", numeric_cols, key="scatter_x")
            y_scatter = st.selectbox("Y Axis", numeric_cols,
                                      index=min(1, len(numeric_cols)-1), key="scatter_y")
            color_col = categorical_cols[0] if categorical_cols else None
            fig_scatter = px.scatter(filtered_df, x=x_scatter, y=y_scatter,
                                      color=color_col,
                                      hover_data=filtered_df.columns.tolist(),
                                      title=f"{x_scatter} vs {y_scatter}")
            fig_scatter.update_layout(height=400, plot_bgcolor='#f8f9fa',
                                       paper_bgcolor='white',
                                       margin=dict(l=0, r=0, t=40, b=0))
            st.plotly_chart(fig_scatter, use_container_width=True)

        with col4:
            st.subheader("🌡️ Correlation Heatmap")
            corr = filtered_df[numeric_cols].corr()
            fig_heat = go.Figure(go.Heatmap(
                z=corr.values, x=corr.columns.tolist(), y=corr.columns.tolist(),
                colorscale='RdYlGn', zmin=-1, zmax=1,
                text=corr.round(2).values, texttemplate='%{text}',
                hovertemplate='<b>%{x} vs %{y}</b><br>Correlation: %{z:.3f}<extra></extra>'
            ))
            fig_heat.update_layout(height=400, paper_bgcolor='white',
                                    margin=dict(l=0, r=0, t=20, b=0),
                                    xaxis=dict(tickfont=dict(size=9)),
                                    yaxis=dict(tickfont=dict(size=9)))
            st.plotly_chart(fig_heat, use_container_width=True)

    st.divider()

    if len(numeric_cols) >= 1:
        st.subheader("📉 Line Chart")
        line_cols = st.multiselect("Select columns to plot", numeric_cols,
                                    default=numeric_cols[:2])
        if line_cols:
            fig_line = go.Figure()
            colors_line = ['#1A6B72', '#E8590C', '#534AB7', '#2ecc71', '#e74c3c']
            for i, col in enumerate(line_cols):
                fig_line.add_trace(go.Scatter(
                    x=list(range(len(filtered_df))), y=filtered_df[col],
                    mode='lines+markers', name=col,
                    line=dict(color=colors_line[i % len(colors_line)], width=2),
                    marker=dict(size=5),
                    hovertemplate=f'<b>{col}</b><br>Value: %{{y:.2f}}<extra></extra>'
                ))
            fig_line.update_layout(height=350, plot_bgcolor='#f8f9fa',
                                    paper_bgcolor='white',
                                    legend=dict(orientation='h', yanchor='bottom', y=1.02),
                                    margin=dict(l=0, r=0, t=40, b=0))
            st.plotly_chart(fig_line, use_container_width=True)

    st.divider()
    st.caption("BeruAnalytics | Interactive Dashboard | Built by Felix Beru Tsinzole")

# ================================================
# PAGE 3 — BERUDATANARRATE
# ================================================
elif page == "🤖 BeruDataNarrate":
    st.markdown("<div class='section-header'>🤖 BeruDataNarrate — AI Report Generator</div>",
                unsafe_allow_html=True)
    st.markdown("Upload your data, describe what it contains, and BeruDataNarrate will generate a professional analysis report with charts, insights and recommendations — powered by AI.")

    uploaded = st.file_uploader("Upload CSV or Excel file", type=['csv', 'xlsx', 'xls'])
    if uploaded:
        df = load_data(uploaded)
        if df is not None:
            st.session_state.df = df
            st.session_state.file_name = uploaded.name
            st.success(f"✅ {uploaded.name} — {len(df):,} rows, {len(df.columns)} columns")
            with st.expander("👀 Preview Data"):
                st.dataframe(df.head(10), use_container_width=True)

    if st.session_state.df is not None:
        df = st.session_state.df
        st.divider()

        col1, col2 = st.columns(2)
        with col1:
            dataset_description = st.text_area(
                "📝 Describe your dataset",
                placeholder="Example: Monthly sales data for our Kenya retail stores...",
                height=120
            )
        with col2:
            user_question = st.text_area(
                "❓ What insights do you need? (optional)",
                placeholder="Example: Which regions are underperforming?",
                height=120
            )

        st.divider()

        # Try secrets first
        secret_key = get_groq_key()
        if secret_key:
            api_key = secret_key
            st.success("🔑 API key loaded automatically")
        else:
            api_key = st.text_input(
                "🔑 Groq API Key",
                type="password",
                placeholder="Enter your Groq API key — get one free at console.groq.com",
                help="Your key is never stored — only used for this session"
            )

        if st.button("🚀 Generate Report", type="primary", use_container_width=True):
            if not dataset_description:
                st.error("Please describe your dataset first")
            elif not api_key:
                st.error("Please enter your Groq API key")
            else:
                with st.spinner("🤖 Analyzing your data with AI..."):
                    try:
                        from groq import Groq
                        import matplotlib
                        matplotlib.use('Agg')
                        import matplotlib.pyplot as plt
                        from docx import Document as DocxDocument
                        from docx.shared import Inches, Pt, RGBColor
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        from reportlab.lib.pagesizes import letter
                        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                        from reportlab.lib.units import inch
                        from reportlab.lib.colors import HexColor, white
                        from reportlab.platypus import SimpleDocTemplate, Paragraph as RLParagraph
                        from reportlab.platypus import Spacer, Image, Table, TableStyle
                        from reportlab.platypus import PageBreak, HRFlowable
                        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

                        groq_client = Groq(api_key=api_key)
                        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()

                        data_summary = {
                            "rows": len(df),
                            "columns": len(df.columns),
                            "column_names": df.columns.tolist(),
                            "sample_data": df.head(5).to_dict(orient='records'),
                            "basic_stats": df[numeric_cols].describe().to_dict() if numeric_cols else {},
                            "missing_values": df.isnull().sum().to_dict(),
                        }

                        prompt = f"""You are a professional data analyst.
DATASET: {dataset_description}
ROWS: {data_summary['rows']}
COLUMNS: {', '.join(data_summary['column_names'])}
SAMPLE: {json.dumps(data_summary['sample_data'], indent=2, default=str)}
STATS: {json.dumps(data_summary['basic_stats'], indent=2, default=str)}
{f"QUESTION: {user_question}" if user_question else ""}

Return ONLY a JSON object:
{{
    "report_title": "title",
    "executive_summary": "2-3 paragraphs",
    "data_overview": "1-2 paragraphs",
    "key_findings": ["finding 1", "finding 2", "finding 3", "finding 4", "finding 5"],
    "detailed_analysis": "3-4 paragraphs",
    "recommendations": ["rec 1", "rec 2", "rec 3"],
    "conclusion": "1-2 paragraphs"
}}
No markdown. No backticks. Only JSON."""

                        response = groq_client.chat.completions.create(
                            model="llama-3.3-70b-versatile",
                            messages=[{"role": "user", "content": prompt}],
                            max_tokens=2000, temperature=0.3
                        )
                        raw = response.choices[0].message.content.strip()
                        if "```json" in raw:
                            raw = raw.split("```json")[1].split("```")[0].strip()
                        elif "```" in raw:
                            raw = raw.split("```")[1].split("```")[0].strip()

                        analysis = json.loads(raw)
                        st.session_state.analysis = analysis
                        st.success("✅ AI analysis complete!")

                        st.subheader(f"📋 {analysis.get('report_title', 'Analysis Report')}")
                        tab1, tab2, tab3, tab4 = st.tabs([
                            "📌 Key Findings", "📊 Analysis",
                            "💡 Recommendations", "📝 Full Report"
                        ])
                        with tab1:
                            for i, f in enumerate(analysis.get('key_findings', []), 1):
                                st.markdown(f"**{i}.** {f}")
                        with tab2:
                            st.markdown("**Executive Summary**")
                            st.write(analysis.get('executive_summary', ''))
                            st.divider()
                            st.markdown("**Detailed Analysis**")
                            st.write(analysis.get('detailed_analysis', ''))
                        with tab3:
                            for i, r in enumerate(analysis.get('recommendations', []), 1):
                                st.markdown(f"**{i}.** {r}")
                        with tab4:
                            st.write(analysis.get('executive_summary', ''))
                            st.write(analysis.get('data_overview', ''))
                            st.write(analysis.get('detailed_analysis', ''))
                            st.write(analysis.get('conclusion', ''))

                    except Exception as e:
                        st.error(f"❌ Analysis failed: {e}")

                with st.spinner("📊 Generating charts..."):
                    try:
                        import matplotlib
                        matplotlib.use('Agg')
                        import matplotlib.pyplot as plt
                        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
                        categorical_cols = df.select_dtypes(include=['object']).columns.tolist()
                        os.makedirs('/tmp/charts', exist_ok=True)
                        chart_paths = []

                        if numeric_cols:
                            fig, ax = plt.subplots(figsize=(10, 5))
                            means = df[numeric_cols].mean()
                            colors_c = plt.cm.Blues(np.linspace(0.4, 0.9, len(means)))
                            ax.bar(range(len(means)), means.values, color=colors_c, edgecolor='white')
                            ax.set_xticks(range(len(means)))
                            ax.set_xticklabels(means.index, rotation=45, ha='right', fontsize=9)
                            ax.set_title('Mean Values by Column', fontweight='bold')
                            ax.set_facecolor('#f8f9fa')
                            ax.grid(axis='y', linestyle=':', alpha=0.5)
                            plt.tight_layout()
                            p1 = '/tmp/charts/chart_1.png'
                            plt.savefig(p1, dpi=150, bbox_inches='tight')
                            plt.close()
                            chart_paths.append(p1)

                        if categorical_cols:
                            cat_col = categorical_cols[0]
                            vc = df[cat_col].value_counts().head(10)
                            fig, ax = plt.subplots(figsize=(10, 5))
                            ax.barh(vc.index.astype(str), vc.values,
                                   color='#1A6B72', edgecolor='white', height=0.6)
                            ax.set_title(f'Distribution of {cat_col}', fontweight='bold')
                            ax.set_facecolor('#f8f9fa')
                            plt.tight_layout()
                            p2 = '/tmp/charts/chart_2.png'
                            plt.savefig(p2, dpi=150, bbox_inches='tight')
                            plt.close()
                            chart_paths.append(p2)

                        if len(numeric_cols) >= 2:
                            corr = df[numeric_cols].corr()
                            fig, ax = plt.subplots(figsize=(8, 6))
                            im = ax.imshow(corr, cmap='RdYlGn', vmin=-1, vmax=1, aspect='auto')
                            ax.set_xticks(range(len(numeric_cols)))
                            ax.set_yticks(range(len(numeric_cols)))
                            ax.set_xticklabels(numeric_cols, rotation=45, ha='right', fontsize=8)
                            ax.set_yticklabels(numeric_cols, fontsize=8)
                            for i in range(len(numeric_cols)):
                                for j in range(len(numeric_cols)):
                                    ax.text(j, i, f'{corr.iloc[i,j]:.2f}',
                                           ha='center', va='center', fontsize=7, fontweight='bold')
                            plt.colorbar(im, ax=ax, shrink=0.8)
                            ax.set_title('Correlation Matrix', fontweight='bold')
                            plt.tight_layout()
                            p3 = '/tmp/charts/chart_3.png'
                            plt.savefig(p3, dpi=150, bbox_inches='tight')
                            plt.close()
                            chart_paths.append(p3)

                        st.success(f"✅ {len(chart_paths)} charts generated!")

                    except Exception as e:
                        st.error(f"Chart error: {e}")
                        chart_paths = []

                if st.session_state.analysis and chart_paths:
                    with st.spinner("📝 Generating Word and PDF reports..."):
                        try:
                            from groq import Groq
                            import matplotlib
                            matplotlib.use('Agg')
                            import matplotlib.pyplot as plt
                            from docx import Document as DocxDocument
                            from docx.shared import Inches, Pt, RGBColor
                            from docx.enum.text import WD_ALIGN_PARAGRAPH
                            from reportlab.lib.pagesizes import letter
                            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                            from reportlab.lib.units import inch
                            from reportlab.lib.colors import HexColor, white
                            from reportlab.platypus import SimpleDocTemplate, Paragraph as RLParagraph
                            from reportlab.platypus import Spacer, Image, PageBreak, HRFlowable
                            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

                            analysis = st.session_state.analysis

                            # WORD
                            word_doc = DocxDocument()
                            section = word_doc.sections[0]
                            section.left_margin = Inches(1)
                            section.right_margin = Inches(1)

                            p = word_doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            r = p.add_run("BeruAnalytics")
                            r.font.size = Pt(28); r.font.bold = True
                            r.font.color.rgb = RGBColor(0x1A, 0x6B, 0x72)

                            p2 = word_doc.add_paragraph()
                            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            r2 = p2.add_run(analysis.get('report_title', 'Report'))
                            r2.font.size = Pt(20); r2.font.bold = True

                            p3 = word_doc.add_paragraph()
                            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            r3 = p3.add_run(
                                f"Generated: {datetime.now().strftime('%B %d, %Y')} | "
                                f"Rows: {len(df):,} | Cols: {len(df.columns)}")
                            r3.font.size = Pt(10)
                            r3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                            word_doc.add_page_break()

                            def add_word_section(title, content_text=None, items=None):
                                h = word_doc.add_paragraph()
                                hr = h.add_run(title.upper())
                                hr.font.size = Pt(13); hr.font.bold = True
                                hr.font.color.rgb = RGBColor(0x1A, 0x6B, 0x72)
                                hr.font.underline = True
                                word_doc.add_paragraph()
                                if content_text:
                                    bp = word_doc.add_paragraph()
                                    br = bp.add_run(content_text)
                                    br.font.size = Pt(11)
                                if items:
                                    for i, item in enumerate(items, 1):
                                        ip = word_doc.add_paragraph()
                                        ir = ip.add_run(f"{i}. {item}")
                                        ir.font.size = Pt(11)
                                word_doc.add_paragraph()

                            add_word_section("Executive Summary", analysis.get('executive_summary'))
                            add_word_section("Data Overview", analysis.get('data_overview'))
                            add_word_section("Key Findings", items=analysis.get('key_findings'))
                            add_word_section("Detailed Analysis", analysis.get('detailed_analysis'))

                            h = word_doc.add_paragraph()
                            hr = h.add_run("VISUALIZATIONS")
                            hr.font.size = Pt(13); hr.font.bold = True
                            hr.font.color.rgb = RGBColor(0x1A, 0x6B, 0x72)
                            hr.font.underline = True

                            for cp in chart_paths:
                                if os.path.exists(cp):
                                    word_doc.add_paragraph()
                                    word_doc.add_picture(cp, width=Inches(5.5))

                            word_doc.add_page_break()
                            add_word_section("Recommendations", items=analysis.get('recommendations'))
                            add_word_section("Conclusion", analysis.get('conclusion'))

                            fp = word_doc.add_paragraph()
                            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            fr = fp.add_run("BeruAnalytics | Powered by AI | Built by Felix Beru Tsinzole")
                            fr.font.size = Pt(9)
                            fr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                            fr.font.italic = True

                            word_path = '/tmp/BeruAnalytics_Report.docx'
                            word_doc.save(word_path)

                            # PDF
                            TEAL = HexColor('#1A6B72')
                            DARK_C = HexColor('#1A1A2E')
                            GRAY_C = HexColor('#555555')

                            pdf_path = '/tmp/BeruAnalytics_Report.pdf'
                            pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter,
                                rightMargin=inch, leftMargin=inch,
                                topMargin=inch, bottomMargin=inch)

                            s_brand = ParagraphStyle('Brand', fontSize=26, textColor=TEAL,
                                alignment=TA_CENTER, fontName='Helvetica-Bold', spaceAfter=8)
                            s_title = ParagraphStyle('Title', fontSize=18, textColor=DARK_C,
                                alignment=TA_CENTER, fontName='Helvetica-Bold', spaceAfter=10)
                            s_sub = ParagraphStyle('Sub', fontSize=10, textColor=GRAY_C,
                                alignment=TA_CENTER, fontName='Helvetica', spaceAfter=6)
                            s_sec = ParagraphStyle('Sec', fontSize=13, textColor=TEAL,
                                fontName='Helvetica-Bold', spaceBefore=14, spaceAfter=6)
                            s_body = ParagraphStyle('Body', fontSize=11, textColor=DARK_C,
                                fontName='Helvetica', spaceAfter=8, alignment=TA_JUSTIFY, leading=16)
                            s_bullet = ParagraphStyle('Bul', fontSize=11, textColor=DARK_C,
                                fontName='Helvetica', spaceAfter=5, leftIndent=20)
                            s_footer = ParagraphStyle('Foot', fontSize=9, textColor=GRAY_C,
                                alignment=TA_CENTER, fontName='Helvetica-Oblique', spaceAfter=4)

                            pc = []
                            pc.append(Spacer(1, 0.5*inch))
                            pc.append(RLParagraph("BeruAnalytics", s_brand))
                            pc.append(Spacer(1, 0.15*inch))
                            pc.append(HRFlowable(width="100%", thickness=2, color=TEAL))
                            pc.append(Spacer(1, 0.2*inch))
                            pc.append(RLParagraph(analysis.get('report_title', 'Report'), s_title))
                            pc.append(RLParagraph(
                                f"Generated: {datetime.now().strftime('%B %d, %Y')} | "
                                f"Rows: {len(df):,} | Columns: {len(df.columns)}", s_sub))
                            pc.append(Spacer(1, 0.2*inch))
                            pc.append(HRFlowable(width="100%", thickness=1, color=GRAY_C))
                            pc.append(PageBreak())

                            def add_pdf_section(title, text=None, items=None):
                                pc.append(RLParagraph(title, s_sec))
                                pc.append(HRFlowable(width="100%", thickness=1, color=TEAL))
                                pc.append(Spacer(1, 0.1*inch))
                                if text:
                                    pc.append(RLParagraph(text, s_body))
                                if items:
                                    for i, item in enumerate(items, 1):
                                        pc.append(RLParagraph(f"{i}. {item}", s_bullet))
                                pc.append(Spacer(1, 0.15*inch))

                            add_pdf_section("EXECUTIVE SUMMARY", analysis.get('executive_summary'))
                            add_pdf_section("DATA OVERVIEW", analysis.get('data_overview'))
                            add_pdf_section("KEY FINDINGS", items=analysis.get('key_findings'))
                            add_pdf_section("DETAILED ANALYSIS", analysis.get('detailed_analysis'))
                            pc.append(PageBreak())

                            pc.append(RLParagraph("VISUALIZATIONS", s_sec))
                            pc.append(HRFlowable(width="100%", thickness=1, color=TEAL))
                            pc.append(Spacer(1, 0.1*inch))
                            for cp in chart_paths:
                                if os.path.exists(cp):
                                    pc.append(Image(cp, width=5.5*inch, height=3.5*inch))
                                    pc.append(Spacer(1, 0.2*inch))

                            pc.append(PageBreak())
                            add_pdf_section("RECOMMENDATIONS", items=analysis.get('recommendations'))
                            add_pdf_section("CONCLUSION", analysis.get('conclusion'))
                            pc.append(HRFlowable(width="100%", thickness=1, color=GRAY_C))
                            pc.append(Spacer(1, 0.1*inch))
                            pc.append(RLParagraph(
                                "BeruAnalytics | Powered by AI | Built by Felix Beru Tsinzole",
                                s_footer))

                            pdf_doc.build(pc)
                            st.success("✅ Reports generated successfully!")

                            st.divider()
                            st.subheader("📥 Download Your Reports")
                            dl1, dl2 = st.columns(2)
                            with dl1:
                                with open(word_path, 'rb') as f:
                                    st.download_button("📝 Download Word Report", f.read(),
                                        file_name="BeruAnalytics_Report.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        use_container_width=True, type="primary")
                            with dl2:
                                with open(pdf_path, 'rb') as f:
                                    st.download_button("📄 Download PDF Report", f.read(),
                                        file_name="BeruAnalytics_Report.pdf",
                                        mime="application/pdf",
                                        use_container_width=True, type="primary")

                        except Exception as e:
                            st.error(f"Report generation error: {e}")
                            import traceback
                            st.code(traceback.format_exc())

# ================================================
# PAGE 4 — AI ASSISTANT
# ================================================
elif page == "💬 AI Assistant":
    st.markdown("<div class='section-header'>💬 AI Data Assistant</div>",
                unsafe_allow_html=True)
    st.markdown("Ask questions about your data in plain English. Get instant AI-powered answers.")

    if st.session_state.df is None:
        st.warning("⚠️ Please upload a file first on the Home page")
        st.stop()

    df = st.session_state.df

    with st.expander("📊 Your Data Summary"):
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Rows", f"{len(df):,}")
        with col2:
            st.metric("Columns", len(df.columns))
        with col3:
            st.metric("Missing Values", df.isnull().sum().sum())
        st.dataframe(df.describe(), use_container_width=True)

    st.divider()

    secret_key = get_groq_key()
    if secret_key:
        api_key_chat = secret_key
        st.success("🔑 API key loaded automatically")
    else:
        api_key_chat = st.text_input("🔑 Groq API Key", type="password",
            placeholder="Enter your Groq API key", key="chat_api_key")

    examples = [
        "What are the top 3 insights from this data?",
        "Which column has the most variation?",
        "Are there any anomalies or outliers I should know about?",
        "What does this data tell us about performance?",
        "Summarize the key statistics in plain English"
    ]

    selected_example = st.selectbox("💡 Try an example question", [""] + examples, key="example_q")
    question = st.text_input("Ask anything about your data:",
                              value=selected_example,
                              placeholder="e.g. Which category has the highest average value?")

    if st.button("🤖 Get Answer", type="primary", use_container_width=True):
        if not question:
            st.error("Please enter a question")
        elif not api_key_chat:
            st.error("Please enter your Groq API key")
        else:
            with st.spinner("🤖 Thinking..."):
                try:
                    from groq import Groq
                    groq_client = Groq(api_key=api_key_chat)
                    numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
                    data_context = {
                        "rows": len(df),
                        "columns": df.columns.tolist(),
                        "sample": df.head(5).to_dict(orient='records'),
                        "stats": df[numeric_cols].describe().to_dict() if numeric_cols else {},
                        "missing": df.isnull().sum().to_dict()
                    }
                    prompt = f"""You are a helpful data analyst assistant.
DATA CONTEXT: {json.dumps(data_context, indent=2, default=str)}
USER QUESTION: {question}
Answer clearly and specifically. Be concise and data-driven. Reference actual numbers."""

                    response = groq_client.chat.completions.create(
                        model="llama-3.3-70b-versatile",
                        messages=[{"role": "user", "content": prompt}],
                        max_tokens=800, temperature=0.3
                    )
                    st.divider()
                    st.markdown("### 🤖 Answer")
                    st.markdown(response.choices[0].message.content.strip())
                except Exception as e:
                    st.error(f"❌ Error: {e}")

# ================================================
# PAGE 5 — DATA EXPLORER
# ================================================
elif page == "📋 Data Explorer":
    st.markdown("<div class='section-header'>📋 Data Explorer</div>",
                unsafe_allow_html=True)

    if st.session_state.df is None:
        st.warning("⚠️ Please upload a file first on the Home page")
        st.stop()

    df = st.session_state.df

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total Rows", f"{len(df):,}")
    with col2:
        st.metric("Total Columns", len(df.columns))
    with col3:
        st.metric("Numeric Columns", len(df.select_dtypes(include=[np.number]).columns))
    with col4:
        st.metric("Missing Values", df.isnull().sum().sum())

    st.divider()

    col_search, col_filter = st.columns([2, 1])
    with col_search:
        search_term = st.text_input("🔍 Search across all columns",
                                     placeholder="Type to filter rows...")
    with col_filter:
        col_select = st.multiselect("Show columns", df.columns.tolist(),
                                     default=df.columns.tolist())

    filtered_df = df[col_select] if col_select else df
    if search_term:
        mask = filtered_df.astype(str).apply(
            lambda x: x.str.contains(search_term, case=False, na=False)
        ).any(axis=1)
        filtered_df = filtered_df[mask]

    st.caption(f"Showing {len(filtered_df):,} of {len(df):,} rows")
    st.dataframe(filtered_df, use_container_width=True, height=400)

    st.divider()
    st.subheader("📥 Export Data")
    exp1, exp2 = st.columns(2)
    with exp1:
        csv_data = filtered_df.to_csv(index=False).encode('utf-8')
        st.download_button("📥 Download as CSV", csv_data,
            file_name="BeruAnalytics_Export.csv", mime="text/csv",
            use_container_width=True, type="primary")
    with exp2:
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            filtered_df.to_excel(writer, index=False, sheet_name='BeruAnalytics')
        st.download_button("📊 Download as Excel", buffer.getvalue(),
            file_name="BeruAnalytics_Export.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True, type="primary")

    st.divider()
    st.subheader("📊 Column Statistics")
    numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    if numeric_cols:
        st.dataframe(df[numeric_cols].describe().T, use_container_width=True)

# ================================================
# PAGE 6 — ABOUT
# ================================================
elif page == "ℹ️ About":
    st.markdown("<div class='section-header'>ℹ️ About BeruAnalytics</div>",
                unsafe_allow_html=True)

    col1, col2 = st.columns([2, 1])
    with col1:
        st.markdown("""
        ### What is BeruAnalytics?
        BeruAnalytics is an AI-powered data analytics platform built to make
        professional data analysis accessible to everyone without needing
        to write code or hire a data analyst.

        Upload your CSV or Excel file and get:
        - **Interactive dashboards** with KPI cards and multiple chart types
        - **AI-generated reports** in Word and PDF format
        - **Natural language insights** — ask questions, get answers
        - **Smart data exploration** with search, filter and export

        ### Who is it for?
        - NGOs and nonprofits needing donor reports
        - Small and medium businesses tracking performance
        - Government departments analyzing public data
        - Consultants and researchers presenting findings
        - Anyone with data who needs professional insights fast

        ### Technology
        - **Python** — data processing
        - **Plotly** — interactive visualizations
        - **Streamlit** — web application framework
        - **Groq LLaMA** — AI analysis and insights
        - **ReportLab + python-docx** — professional report generation
        """)

    with col2:
        st.markdown("""
        ### Built By
        **Felix Beru Tsinzole**
        Data Analyst | Nairobi, Kenya 🇰🇪

        With expertise in:
        - Public health data analysis
        - Financial data analytics
        - BI dashboard development
        - AI-powered data tools
        - Data Cleaning and Transformation
        """)
        st.divider()
        st.markdown("**📊 Portfolio:**")
        st.markdown("[LinkedIn](http://www.linkedin.com/in/felix-beru-04b905280)")
        st.markdown("[GitHub](https://github.com/FelixBeruTheAnalyst)")
        st.markdown("[Maternal Health Dashboard](https://kenya-maternal-health-anc-analysis-oasrbxen82vgwacmq4vbny.streamlit.app/)")
        st.markdown("[CPI Inflation Dashboard](https://kenya-cpi-inflation-analysis-l28f5kfd6qypapa6jrkazq.streamlit.app/)")
        st.divider()
        st.markdown("**📧 Contact:**")
        st.markdown("berufelix42@gmail.com")
        st.markdown("079-797-1074")

    st.divider()

    st.markdown("<div class='section-header'>🔒 Data Privacy & Security</div>",
                unsafe_allow_html=True)

    col3, col4 = st.columns(2)
    with col3:
        st.success("""
        **✅ What BeruAnalytics does NOT do:**
        - Does not store your data permanently
        - Does not share your data with other users
        - Does not write your data to any database
        - Each session is completely isolated

        All uploaded data lives only in memory for the
        duration of your browser session.
        When you close the tab — it is gone.
        """)
    with col4:
        st.warning("""
        **⚠️ What you should know:**
        - A data summary (first 5 rows + statistics)
          is sent to Groq AI for analysis
        - Data is transmitted over HTTPS (encrypted)
        - Streamlit Cloud temporarily processes
          your data during the session

        Your data is never permanently stored —
        but a summary is shared with our AI provider.
        """)

    st.divider()

    col5, col6 = st.columns(2)
    with col5:
        st.success("""
        **✅ Safe to Upload**
        - Aggregated business data
        - Public datasets
        - Anonymized survey results
        - Financial summaries without account numbers
        - Sales and operational data
        - Research and monitoring data
        """)
    with col6:
        st.error("""
        **❌ Avoid Uploading**
        - Patient records with names or IDs
        - Employee personal information
        - Financial data with account numbers
        - Any data covered by GDPR
        - Kenya Data Protection Act regulated data
        - Passwords or authentication credentials
        """)

    st.divider()

    st.info("""
    **📋 Our Privacy Commitment**

    BeruAnalytics does not store your data. All uploads are processed
    in memory and deleted when your session ends. A data summary is
    sent to our AI provider (Groq) for analysis purposes only.

    **Do not upload personally identifiable or confidential information.**

    For enterprise deployments requiring full data privacy,
    contact us about on-premise installation options.
    """)

    st.divider()
    st.caption("BeruAnalytics v1.0 | Built by Felix Beru Tsinzole | Nairobi, Kenya 🇰🇪")
